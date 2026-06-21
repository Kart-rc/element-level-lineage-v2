import importlib.util
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
BUILDER_PATH = ROOT / "scripts" / "build_prd_docx.py"


def load_builder():
    assert BUILDER_PATH.exists(), "DOCX builder module is missing"
    spec = importlib.util.spec_from_file_location("build_prd_docx", BUILDER_PATH)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class MarkdownParserTests(unittest.TestCase):
    def test_parse_blocks_preserves_headings_lists_and_tables(self):
        builder = load_builder()
        source = """# Title

Intro with **bold** and `code`.

## Section

1. First
2. Second

- Alpha
- Beta

| ID | Requirement |
|---|---|
| A-1 | Works |
"""

        blocks = builder.parse_markdown(source)

        self.assertEqual(
            [block.kind for block in blocks],
            ["heading", "paragraph", "heading", "ordered_list", "bullet_list", "table"],
        )
        self.assertEqual(blocks[0].level, 1)
        self.assertEqual(blocks[3].items, ["First", "Second"])
        self.assertEqual(blocks[5].rows, [["ID", "Requirement"], ["A-1", "Works"]])


class DocxBuilderTests(unittest.TestCase):
    def test_build_docx_applies_business_brief_structure(self):
        builder = load_builder()
        source = """# Product Requirements Document: Test Product

**Enterprise Data Lineage and Impact Analysis Platform**  
**Status:** Draft  
**Date:** June 20, 2026

## Executive Summary

This is the executive summary.

## Requirements

| ID | Requirement | Acceptance criteria |
|---|---|---|
| T-001 | Render a table | Table geometry is fixed |

A source is supported when:

1. First numbered item
2. Second numbered item
"""

        with tempfile.TemporaryDirectory() as tmp:
            source_path = Path(tmp) / "source.md"
            output_path = Path(tmp) / "output.docx"
            source_path.write_text(source, encoding="utf-8")

            builder.build_docx(source_path, output_path)

            self.assertTrue(output_path.exists())
            document = Document(output_path)
            self.assertEqual(len(document.sections), 1)
            self.assertAlmostEqual(document.sections[0].left_margin.inches, 1.0, places=2)
            self.assertEqual(document.styles["Normal"].font.name, "Calibri")
            title_properties = document.styles["Title"]._element.pPr
            self.assertIsNone(title_properties.find(qn("w:pBdr")))
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(document.tables[0].autofit, False)
            table_width = document.tables[0]._tbl.tblPr.find(qn("w:tblW"))
            self.assertEqual(table_width.get(qn("w:type")), "dxa")
            self.assertEqual(table_width.get(qn("w:w")), "9360")
            numbered = [p for p in document.paragraphs if p._p.pPr is not None and p._p.pPr.numPr is not None]
            self.assertEqual(len(numbered), 2)
            self.assertTrue(numbered[0].paragraph_format.keep_with_next)
            lead_in = next(p for p in document.paragraphs if p.text == "A source is supported when:")
            self.assertTrue(lead_in.paragraph_format.keep_with_next)
            all_text = "\n".join(p.text for p in document.paragraphs)
            self.assertIn("Test Product", all_text)
            self.assertIn("Executive Summary", all_text)
            self.assertEqual(all_text.count("Enterprise Data Lineage and Impact Analysis Platform"), 1)


if __name__ == "__main__":
    unittest.main()
