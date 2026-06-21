#!/usr/bin/env python3
"""Build the Throughline PRD DOCX from its canonical Markdown source."""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor, Twips


PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

BLACK = RGBColor(0x22, 0x22, 0x22)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x6A, 0x73)
LIGHT_GRAY = "F2F4F7"
BORDER_GRAY = "D9DEE5"
LINK_BLUE = "0563C1"


class Block:
    def __init__(
        self,
        kind: str,
        *,
        level: Optional[int] = None,
        text: Optional[str] = None,
        items: Optional[List[str]] = None,
        rows: Optional[List[List[str]]] = None,
    ) -> None:
        self.kind = kind
        self.level = level
        self.text = text
        self.items = items or []
        self.rows = rows or []


HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.+?)\s*$")
BULLET_RE = re.compile(r"^-\s+(.+?)\s*$")
TABLE_SEPARATOR_RE = re.compile(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$")
INLINE_RE = re.compile(r"(\[[^\]]+\]\(https?://[^)]+\)|\*\*[^*]+\*\*|`[^`]+`)")


def _join_paragraph_lines(lines: List[str]) -> str:
    result = ""
    for index, line in enumerate(lines):
        hard_break = line.endswith("  ")
        result += line.rstrip()
        if index < len(lines) - 1:
            result += "\n" if hard_break else " "
    return result


def _parse_table_row(line: str) -> List[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def parse_markdown(source: str) -> List[Block]:
    """Parse the controlled Markdown subset used by the PRD."""
    lines = source.splitlines()
    blocks: List[Block] = []
    index = 0

    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue

        heading = HEADING_RE.match(line)
        if heading:
            blocks.append(Block("heading", level=len(heading.group(1)), text=heading.group(2)))
            index += 1
            continue

        if (
            line.lstrip().startswith("|")
            and index + 1 < len(lines)
            and TABLE_SEPARATOR_RE.match(lines[index + 1].strip())
        ):
            rows = [_parse_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                rows.append(_parse_table_row(lines[index]))
                index += 1
            blocks.append(Block("table", rows=rows))
            continue

        ordered = ORDERED_RE.match(line)
        if ordered:
            items = []
            while index < len(lines):
                match = ORDERED_RE.match(lines[index])
                if not match:
                    break
                items.append(match.group(1))
                index += 1
            blocks.append(Block("ordered_list", items=items))
            continue

        bullet = BULLET_RE.match(line)
        if bullet:
            items = []
            while index < len(lines):
                match = BULLET_RE.match(lines[index])
                if not match:
                    break
                items.append(match.group(1))
                index += 1
            blocks.append(Block("bullet_list", items=items))
            continue

        paragraph_lines = []
        while index < len(lines) and lines[index].strip():
            current = lines[index]
            if paragraph_lines and (
                HEADING_RE.match(current)
                or ORDERED_RE.match(current)
                or BULLET_RE.match(current)
                or (
                    current.lstrip().startswith("|")
                    and index + 1 < len(lines)
                    and TABLE_SEPARATOR_RE.match(lines[index + 1].strip())
                )
            ):
                break
            paragraph_lines.append(current)
            index += 1
        blocks.append(Block("paragraph", text=_join_paragraph_lines(paragraph_lines)))

    return blocks


def _set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: Optional[float] = None,
    color: Optional[RGBColor] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style, name: str, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def _configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    _set_style_font(normal, "Calibri", 11, BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    _set_style_font(title, "Calibri", 23, BLACK, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.0
    title.paragraph_format.keep_with_next = True
    _remove_children(title._element.get_or_add_pPr(), "w:pBdr")

    subtitle = styles["Subtitle"]
    _set_style_font(subtitle, "Calibri", 14, MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.0
    subtitle.paragraph_format.keep_with_next = True

    heading_tokens = [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]
    for name, size, color, before, after in heading_tokens:
        style = styles[name]
        _set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    for style_name, size, color, bold in [
        ("PRD Kicker", 9, BLUE, True),
        ("PRD Metadata", 10, BLACK, False),
        ("PRD Table", 8.5, BLACK, False),
        ("PRD Table Header", 8.5, BLACK, True),
    ]:
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[style_name]
        _set_style_font(style, "Calibri", size, color, bold=bold)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(2 if "Table" in style_name else 3)
        style.paragraph_format.line_spacing = 1.05 if "Table" in style_name else 1.0
        style.paragraph_format.widow_control = True


def _append_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    _set_run_font(run, size=8.5, color=MUTED)


def _configure_section(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION.NEW_PAGE

    header_p = section.header.paragraphs[0]
    header_p.clear()
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header_p.add_run("Throughline | Product Requirements Document")
    _set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = header_p.add_run("\tOpenLineage Platform")
    _set_run_font(right, size=8.5, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.clear()
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = footer_p.add_run("Throughline | June 20, 2026")
    _set_run_font(left, size=8.5, color=MUTED)
    page_label = footer_p.add_run("\tPage ")
    _set_run_font(page_label, size=8.5, color=MUTED)
    _append_page_field(footer_p)


def _add_hyperlink(paragraph, label: str, url: str) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([run_properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_plain_run(paragraph, text: str) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    _set_run_font(run, size=11, color=BLACK)


def add_rich_text(paragraph, text: str, *, size: float = 11) -> None:
    """Add supported inline Markdown to a Word paragraph."""
    cursor = 0
    for match in INLINE_RE.finditer(text):
        plain = text[cursor : match.start()]
        _add_plain_run(paragraph, plain)
        token = match.group(0)
        if token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            _add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, size=size, color=BLACK, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, name="Consolas", size=max(8.5, size - 1), color=DARK_BLUE)
        cursor = match.end()
    _add_plain_run(paragraph, text[cursor:])

    for run in paragraph.runs:
        if run.font.size is None or run.font.size.pt == 11:
            run.font.size = Pt(size)


def _add_numbering_definition(document: Document, *, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(element.get(qn("w:numId"))) for element in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "decimal" if ordered else "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "720")
    indentation.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.extend([tabs, indentation, spacing])

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    run_properties.append(fonts)

    level.extend([start, num_format, level_text, justification, suffix, paragraph_properties, run_properties])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    num_properties = paragraph_properties.get_or_add_numPr()
    level = num_properties.get_or_add_ilvl()
    level.set(qn("w:val"), "0")
    number = num_properties.get_or_add_numId()
    number.set(qn("w:val"), str(num_id))
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.paragraph_format.widow_control = True


def _remove_children(parent, tag: str) -> None:
    for child in list(parent.findall(qn(tag))):
        parent.remove(child)


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _table_widths(rows: List[List[str]]) -> List[int]:
    count = len(rows[0])
    header = [cell.lower() for cell in rows[0]]
    patterns = {
        ("persona", "primary job", "required outcome"): [2016, 3312, 4032],
        ("id", "requirement", "acceptance criteria"): [1440, 3600, 4320],
        ("entity", "system of record", "notes"): [1872, 2592, 4896],
        ("severity", "meaning", "example"): [1440, 3168, 4752],
        ("risk", "impact", "mitigation"): [2160, 1872, 5328],
        ("requirement area", "primary sections", "launch evidence"): [2880, 2160, 4320],
        ("dependency", "required owner"): [5040, 4320],
        ("category", "requirement"): [2160, 7200],
    }
    key = tuple(header)
    if key in patterns:
        return patterns[key]
    if count == 2:
        return [2808, 6552]
    if count == 3:
        return [1872, 3744, 3744]
    base = CONTENT_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _set_cell_margins(cell) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    _remove_children(cell_properties, "w:tcMar")
    margins = OxmlElement("w:tcMar")
    for side, value in [("top", 80), ("bottom", 80), ("start", 120), ("end", 120)]:
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    cell_properties.append(margins)


def _configure_table_geometry(table, widths: List[int]) -> None:
    table.autofit = False
    table_properties = table._tbl.tblPr
    table_width = _ensure_child(table_properties, "w:tblW")
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))

    _remove_children(table_properties, "w:tblInd")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    table_properties.append(indent)

    _remove_children(table_properties, "w:tblLayout")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_properties.append(layout)

    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for grid_column, width in zip(grid_columns, widths):
        grid_column.set(qn("w:w"), str(width))

    for row in table.rows:
        row_properties = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_properties.append(cant_split)
        for cell, width in zip(row.cells, widths):
            cell.width = Twips(width)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = _ensure_child(cell_properties, "w:tcW")
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(width))
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _mark_repeating_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    _remove_children(properties, "w:shd")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _add_table(document: Document, rows: List[List[str]]) -> None:
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    widths = _table_widths(rows)
    _configure_table_geometry(table, widths)
    _mark_repeating_header(table.rows[0])

    for row_index, source_row in enumerate(rows):
        for column_index, value in enumerate(source_row):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.style = "PRD Table Header" if row_index == 0 else "PRD Table"
            paragraph.paragraph_format.keep_together = True
            add_rich_text(paragraph, value, size=8.5)
            if row_index == 0:
                _shade_cell(cell, LIGHT_GRAY)
                for run in paragraph.runs:
                    run.bold = True
            else:
                _shade_cell(cell, "FFFFFF")

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)


def _add_masthead(document: Document, blocks: List[Block]) -> int:
    title_block = blocks[0]
    full_title = title_block.text or "Product Requirements Document"
    product_name = full_title.split(":", 1)[1].strip() if ":" in full_title else "Throughline"

    kicker = document.add_paragraph(style="PRD Kicker")
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    _set_run_font(run, size=9, color=BLUE, bold=True)

    title = document.add_paragraph(style="Title")
    _remove_children(title._p.get_or_add_pPr(), "w:pBdr")
    title.add_run(product_name)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("Enterprise Data Lineage and Impact Analysis Platform")

    consumed = 1
    if len(blocks) > 1 and blocks[1].kind == "paragraph":
        for line in (blocks[1].text or "").split("\n"):
            plain_line = line.replace("**", "").replace("`", "").strip()
            if plain_line == "Enterprise Data Lineage and Impact Analysis Platform":
                continue
            metadata = document.add_paragraph(style="PRD Metadata")
            add_rich_text(metadata, line, size=10)
        consumed += 1

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return consumed


def _render_blocks(document: Document, blocks: List[Block], start: int) -> None:
    remaining = blocks[start:]
    for block_index, block in enumerate(remaining):
        next_block = remaining[block_index + 1] if block_index + 1 < len(remaining) else None
        if block.kind == "heading":
            style_index = min(max((block.level or 2) - 1, 1), 3)
            paragraph = document.add_paragraph(style=f"Heading {style_index}")
            add_rich_text(paragraph, block.text or "", size=[0, 16, 13, 12][style_index])
        elif block.kind == "paragraph":
            paragraph = document.add_paragraph(style="Normal")
            if next_block is not None and next_block.kind in {"ordered_list", "bullet_list", "table"}:
                paragraph.paragraph_format.keep_with_next = True
            for part_index, part in enumerate((block.text or "").split("\n")):
                if part_index:
                    paragraph.add_run().add_break(WD_BREAK.LINE)
                add_rich_text(paragraph, part, size=11)
        elif block.kind in {"ordered_list", "bullet_list"}:
            ordered = block.kind == "ordered_list"
            num_id = _add_numbering_definition(document, ordered=ordered)
            for item_index, item in enumerate(block.items):
                paragraph = document.add_paragraph(style="Normal")
                _apply_numbering(paragraph, num_id)
                if item_index == len(block.items) - 2:
                    paragraph.paragraph_format.keep_with_next = True
                add_rich_text(paragraph, item, size=11)
        elif block.kind == "table":
            _add_table(document, block.rows)


def build_docx(source_path: Path, output_path: Path) -> None:
    source = source_path.read_text(encoding="utf-8")
    blocks = parse_markdown(source)
    if not blocks or blocks[0].kind != "heading" or blocks[0].level != 1:
        raise ValueError("The PRD source must begin with a level-one title")

    document = Document()
    _configure_section(document)
    _configure_styles(document)
    document.core_properties.title = "Throughline - Enterprise Data Lineage and Impact Analysis Platform PRD"
    document.core_properties.subject = "Production product requirements with OpenLineage as the canonical lineage format"
    document.core_properties.author = "Throughline Product Team"
    document.core_properties.keywords = "Throughline, OpenLineage, data lineage, impact analysis, PRD"

    consumed = _add_masthead(document, blocks)
    _render_blocks(document, blocks, consumed)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path, help="Canonical Markdown PRD")
    parser.add_argument("output", type=Path, help="Output DOCX path")
    args = parser.parse_args()
    build_docx(args.source, args.output)


if __name__ == "__main__":
    main()
